#!/usr/bin/env python3
"""Synthetic DOCX fixtures for the Phase 13.3 ingestion tests.

Every fixture here is generated programmatically so the test suite never has to
commit a real, private, or copyrighted Word document. Nothing downloads
anything and no business content is used.

The bilingual fixture deliberately mixes Traditional Chinese and English to
prove Unicode fidelity, and repeats a heading name under two different parents
to prove heading paths stay distinct.
"""

from __future__ import annotations

import io
import zipfile


def build_structured_docx() -> bytes:
    """Headings, paragraphs, a table, and paragraph/table interleaving.

    Document order is deliberately: heading, paragraph, heading, paragraph,
    table, paragraph, heading, heading, paragraph -- so that a extractor which
    emitted "all paragraphs then all tables" would visibly fail ordering tests.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Governance", level=1)
    doc.add_paragraph("Synthetic governance overview.")

    doc.add_heading("Risk Controls", level=2)
    doc.add_paragraph("\u6cbb\u7406\u67b6\u69cb\u8207\u98a8\u96aa\u63a7\u5236\u3002")

    table = doc.add_table(rows=3, cols=2)
    cells = [
        ("Control", "Owner"),
        ("Access", "Security"),
        ("Review", "Governance"),
    ]
    for row_index, (left, right) in enumerate(cells):
        table.cell(row_index, 0).text = left
        table.cell(row_index, 1).text = right

    doc.add_paragraph("Paragraph after the table.")

    doc.add_heading("Testing", level=1)
    doc.add_heading("Risk Controls", level=2)
    doc.add_paragraph("Different repeated-heading section.")

    return _to_bytes(doc)


def build_bilingual_docx() -> bytes:
    """Minimal bilingual document."""
    from docx import Document

    doc = Document()
    doc.add_heading("\u6e2c\u8a66\u6587\u4ef6", level=1)
    doc.add_paragraph("Governance and Risk Controls")
    doc.add_paragraph("\u6cbb\u7406\u67b6\u69cb\u8207\u98a8\u96aa\u63a7\u5236\u3002")
    return _to_bytes(doc)


def build_empty_docx() -> bytes:
    """A structurally valid DOCX with no meaningful text at all."""
    from docx import Document

    doc = Document()
    # Layout-only empty paragraphs: whitespace and empty runs, no real text.
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("\t")
    return _to_bytes(doc)


def build_unstyled_bold_docx() -> bytes:
    """Bold text with no heading style must stay ordinary paragraph content."""
    from docx import Document

    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Looks Like A Heading")
    run.bold = True
    run.font.size = None
    doc.add_paragraph("Body text follows.")
    return _to_bytes(doc)


def build_image_only_docx() -> bytes:
    """A DOCX whose only content is an embedded image, with no body text.

    Used to prove that image content is never OCR'd and that a document with no
    extractable text fails deterministically instead of succeeding empty.
    """
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_picture(_tiny_png(), width=Inches(1))
    return _to_bytes(doc)


def build_table_only_docx() -> bytes:
    """Table content with no body paragraphs, to prove tables are not dropped."""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "\u9805\u76ee"
    table.cell(1, 0).text = "Value"
    table.cell(1, 1).text = "\u6578\u503c"
    return _to_bytes(doc)


def build_ordinary_zip() -> bytes:
    """A plain ZIP archive that is not an OOXML package at all."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("hello.txt", "just a zip, not a docx")
    return buffer.getvalue()


def build_xlsx_like_package() -> bytes:
    """An OOXML package that is a spreadsheet, not a word-processing document."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("xl/workbook.xml", "<workbook/>")
    return buffer.getvalue()


def build_corrupt_docx() -> bytes:
    """Starts with a ZIP signature but the container is broken."""
    return b"PK\x03\x04" + b"\x00\xff broken ooxml container \x01\x02\x03"


def build_encrypted_like_docx() -> bytes:
    """An OLE/CFB container, which is how password-protected DOCX is wrapped."""
    return b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512


def _to_bytes(doc) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _tiny_png() -> io.BytesIO:
    """A 1x1 PNG, built inline so no binary fixture is committed."""
    import struct
    import zlib

    def chunk(kind: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + kind
            + data
            + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF)
        )

    header = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    pixels = zlib.compress(b"\x00\xff\xff\xff")
    png = (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", header)
        + chunk(b"IDAT", pixels)
        + chunk(b"IEND", b"")
    )
    return io.BytesIO(png)
