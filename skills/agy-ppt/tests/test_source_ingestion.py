#!/usr/bin/env python3
"""Phase 13.1/13.2 -- deterministic tests for source ingestion & locators.

Every test here is deterministic: no real Codex/Kiro process is launched, no
subscription quota is consumed, no network request is made, and no real
third-party or confidential document is used. All PDF fixtures are synthesised
byte-by-byte by ``tests/helpers/synthetic_pdf.py``, and all Markdown/text
fixtures are written inline.

These tests cover extraction only. They never assert that an extracted block is
semantically important, nor that it constitutes a Phase 12 source unit --
promoting a block to a semantic source unit is an AGY decision that this layer
deliberately cannot make.
"""

from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parents[1] / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

HELPERS_DIR = Path(__file__).resolve().parent / "helpers"
if str(HELPERS_DIR) not in sys.path:
    sys.path.insert(0, str(HELPERS_DIR))

from synthetic_pdf import (  # noqa: E402
    build_corrupt_pdf,
    build_text_pdf,
    build_textless_pdf,
)
from synthetic_docx import (  # noqa: E402
    build_bilingual_docx,
    build_corrupt_docx,
    build_empty_docx,
    build_encrypted_like_docx,
    build_image_only_docx,
    build_ordinary_zip,
    build_structured_docx,
    build_table_only_docx,
    build_unstyled_bold_docx,
    build_xlsx_like_package,
)

import source_ingestion as si  # noqa: E402
from source_grounding import (  # noqa: E402
    SourceInventory,
    compute_source_digest,
    validate_locator,
)

PAGE_TEXTS = [
    "Synthetic Governance Overview",
    "Synthetic Risk Controls",
    "Synthetic Appendix",
]

MARKDOWN_FIXTURE = """# Overview

Intro paragraph.

## Architecture

Architecture text.

### Validation

Validation text.

# Testing

## Architecture

Different architecture section.

```text
# Fake Heading
```
"""

TEXT_FIXTURE = (
    "\u7b2c\u4e00\u6bb5\u5167\u5bb9\u3002\n"
    "Second line of paragraph one.\n"
    "\n"
    "\u7b2c\u4e8c\u6bb5\u5167\u5bb9\u3002\n"
    "\n"
    "\n"
    "Third paragraph here.\n"
)

#: Synthetic HTML only. No real website snapshot, and every external reference
#: points at the reserved ``.invalid`` TLD so it can never resolve.
HTML_FIXTURE = """<!DOCTYPE html>
<html><head><title>Ignored Title</title>
<style>.hidden{display:none}</style>
<script>var a = 1;</script>
<script type="application/ld+json">{"@type":"Article","headline":"JSONLD Headline"}</script>
</head>
<body>
<h1>Architecture</h1>
<p>Intro with <strong>bold</strong>, <em>emphasis</em> and
<a href="https://example.invalid/page">Example Link</a>.</p>
<!-- this comment must never be ingested -->
<h2>Overview</h2>
<p>\u6cbb\u7406\u67b6\u69cb and Governance Architecture.</p>
<ul><li>alpha</li><li>beta<ul><li>beta-nested</li></ul></li></ul>
<ol><li>first</li><li>second</li></ol>
<table>
<tr><th>Control</th><th>Owner</th></tr>
<tr><td rowspan="2">Access</td><td>Security</td></tr>
<tr><td>Review</td></tr>
<tr><td colspan="2">Wide cell</td></tr>
</table>
<p>After table</p>
<noscript>noscript content</noscript>
<iframe src="https://example.invalid/frame"></iframe>
<img src="https://example.invalid/image.png">
<script src="https://example.invalid/app.js"></script>
<h1>Testing</h1>
<h3>Deep Jump</h3>
<p>tail</p>
<h2>Overview</h2>
<p>second overview</p>
</body></html>
"""


class IngestionTestCase(unittest.TestCase):
    """Shared temporary workspace; every fixture is synthetic."""

    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.root = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def write(self, name: str, data: bytes | str) -> Path:
        path = self.root / name
        if isinstance(data, str):
            path.write_text(data, encoding="utf-8")
        else:
            path.write_bytes(data)
        return path

    def pdf(self, name: str = "doc.pdf") -> Path:
        return self.write(name, build_text_pdf(PAGE_TEXTS))

    def markdown(self, name: str = "notes.md") -> Path:
        return self.write(name, MARKDOWN_FIXTURE)

    def text(self, name: str = "plain.txt") -> Path:
        return self.write(name, TEXT_FIXTURE)

    def docx(self, name: str = "structured.docx") -> Path:
        return self.write(name, build_structured_docx())

    def html(self, name: str = "doc.html") -> Path:
        return self.write(name, HTML_FIXTURE)


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------
class TestFormatDetection(IngestionTestCase):
    def test_pdf_detected(self) -> None:
        self.assertEqual(si.detect_source_format(self.pdf()), si.FORMAT_PDF)

    def test_markdown_detected(self) -> None:
        self.assertEqual(si.detect_source_format(self.markdown()), si.FORMAT_MARKDOWN)

    def test_text_detected(self) -> None:
        self.assertEqual(si.detect_source_format(self.text()), si.FORMAT_TEXT)

    def test_unsupported_binary_rejected(self) -> None:
        png = self.write("image.png", b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR")
        self.assertEqual(si.detect_source_format(png), si.FORMAT_UNSUPPORTED)
        with self.assertRaises(si.SourceFormatUnsupported) as ctx:
            si.ingest_source(png, "src_png")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_FORMAT_UNSUPPORTED)

    def test_pdf_signature_wins_over_extension(self) -> None:
        """A PDF named .txt is still routed to the PDF extractor."""
        mislabelled = self.write("actually.txt", build_text_pdf(PAGE_TEXTS))
        self.assertEqual(si.detect_source_format(mislabelled), si.FORMAT_PDF)

    def test_pdf_extension_without_signature_is_unsupported(self) -> None:
        fake = self.write("fake.pdf", b"this is not a pdf at all")
        self.assertEqual(si.detect_source_format(fake), si.FORMAT_UNSUPPORTED)

    def test_missing_file_raises(self) -> None:
        with self.assertRaises(si.SourceFileNotFound) as ctx:
            si.detect_source_format(self.root / "absent.md")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_FILE_NOT_FOUND)

    def test_invalid_source_id_rejected(self) -> None:
        with self.assertRaises(si.SourceIngestionError):
            si.ingest_source(self.markdown(), "not-a-valid-id")

    def test_docx_detected(self) -> None:
        self.assertEqual(si.detect_source_format(self.docx()), si.FORMAT_DOCX)

    def test_ordinary_zip_not_misclassified_as_docx(self) -> None:
        """A ZIP signature alone must not make something a DOCX."""
        plain = self.write("archive.zip", build_ordinary_zip())
        self.assertEqual(si.detect_source_format(plain), si.FORMAT_UNSUPPORTED)
        with self.assertRaises(si.SourceFormatUnsupported):
            si.ingest_source(plain, "src_zip")

    def test_other_ooxml_package_not_misclassified_as_docx(self) -> None:
        sheet = self.write("book.xlsx", build_xlsx_like_package())
        self.assertEqual(si.detect_source_format(sheet), si.FORMAT_UNSUPPORTED)

    def test_docx_extension_without_zip_signature_routes_to_docx(self) -> None:
        """Encrypted DOCX is OLE-wrapped, so it still gets a DOCX diagnostic."""
        enc = self.write("locked.docx", build_encrypted_like_docx())
        self.assertEqual(si.detect_source_format(enc), si.FORMAT_DOCX)

    def test_html_detected(self) -> None:
        self.assertEqual(si.detect_source_format(self.html()), si.FORMAT_HTML)

    def test_htm_extension_detected(self) -> None:
        path = self.write("page.htm", HTML_FIXTURE)
        self.assertEqual(si.detect_source_format(path), si.FORMAT_HTML)

    def test_text_with_angle_brackets_is_not_html(self) -> None:
        """A .txt file mentioning <div> must stay plain text."""
        path = self.write(
            "notes.txt", "Use a <div> and a <p> tag in your markup.\nSecond line.\n"
        )
        self.assertEqual(si.detect_source_format(path), si.FORMAT_TEXT)
        blocks = si.ingest_source(path, "src_txtangle").blocks
        self.assertEqual(blocks[0].block_type, si.BLOCK_PARAGRAPH)

    def test_xml_is_not_detected_as_html(self) -> None:
        path = self.write("data.xml", '<?xml version="1.0"?><root><item>x</item></root>\n')
        self.assertEqual(si.detect_source_format(path), si.FORMAT_UNSUPPORTED)

    def test_html_extension_without_markup_is_unsupported(self) -> None:
        """Content clearly contradicting the extension is not trusted."""
        path = self.write("prose.html", "just prose, no markup at all\n")
        self.assertEqual(si.detect_source_format(path), si.FORMAT_UNSUPPORTED)

    def test_extensionless_html_detected_by_marker(self) -> None:
        path = self.write("page", "<!DOCTYPE html><html><body><p>hi</p></body></html>\n")
        self.assertEqual(si.detect_source_format(path), si.FORMAT_HTML)

    def test_empty_file_is_unsupported(self) -> None:
        path = self.write("zero.html", "")
        self.assertEqual(si.detect_source_format(path), si.FORMAT_UNSUPPORTED)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
class TestPdfExtraction(IngestionTestCase):
    def test_multi_page_extraction(self) -> None:
        result = si.ingest_source(self.pdf(), "src_pdf")
        self.assertEqual(result.source_format, si.FORMAT_PDF)
        self.assertEqual(result.block_count, 3)
        self.assertEqual([b.block_type for b in result.blocks], [si.BLOCK_PAGE] * 3)
        for block, expected in zip(result.blocks, PAGE_TEXTS):
            self.assertIn(expected, block.text)

    def test_page_locators_are_one_based(self) -> None:
        """Page 1 text must land on locator page 1, with no off-by-one."""
        result = si.ingest_source(self.pdf(), "src_pdf")
        pages = [b.locator["start"] for b in result.blocks]
        self.assertEqual(pages, [1, 2, 3])
        for block in result.blocks:
            page = block.locator["start"]
            self.assertEqual(block.locator["end"], page)
            self.assertEqual(block.metadata["page"], page)
            # The text on page N must be the fixture text authored for page N.
            self.assertIn(PAGE_TEXTS[page - 1], block.text)
        self.assertEqual(min(pages), 1, "page numbering must start at 1, not 0")

    def test_deterministic_block_ids(self) -> None:
        result = si.ingest_source(self.pdf(), "src_pdf")
        for block in result.blocks:
            self.assertEqual(
                block.block_id,
                si.compute_block_id("src_pdf", block.locator, block.ordinal),
            )
            self.assertTrue(block.block_id.startswith("blk:src_pdf:"))
        self.assertEqual(
            len({b.block_id for b in result.blocks}), result.block_count
        )

    def test_repeated_ingestion_is_identical(self) -> None:
        path = self.pdf()
        first = si.ingest_source(path, "src_pdf").to_dict()
        second = si.ingest_source(path, "src_pdf").to_dict()
        self.assertEqual(first, second)

    def test_block_ordering_is_deterministic(self) -> None:
        result = si.ingest_source(self.pdf(), "src_pdf")
        ordinals = [b.ordinal for b in result.blocks]
        self.assertEqual(ordinals, sorted(ordinals))
        self.assertEqual(ordinals, list(range(1, result.block_count + 1)))

    def test_textless_pdf_fails_clearly(self) -> None:
        """A scanned/image-only PDF must fail, never succeed with empty text."""
        scan = self.write("scan.pdf", build_textless_pdf(2))
        self.assertEqual(si.detect_source_format(scan), si.FORMAT_PDF)
        with self.assertRaises(si.SourceTextUnavailable) as ctx:
            si.ingest_source(scan, "src_scan")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_TEXT_UNAVAILABLE)
        self.assertIn("OCR is not supported", str(ctx.exception))

    def test_corrupted_pdf_fails_as_extraction_failure(self) -> None:
        """Structurally broken is distinct from validly textless."""
        broken = self.write("broken.pdf", build_corrupt_pdf())
        with self.assertRaises(si.SourceExtractionFailed) as ctx:
            si.ingest_source(broken, "src_broken")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_EXTRACTION_FAILED)
        self.assertNotEqual(
            ctx.exception.error_code,
            si.ERROR_SOURCE_TEXT_UNAVAILABLE,
            "a corrupt PDF must not be reported as a valid textless PDF",
        )

    def test_no_ocr_dependency_is_imported(self) -> None:
        si.ingest_source(self.pdf(), "src_pdf")
        forbidden = {"pytesseract", "tesserocr", "easyocr", "paddleocr"}
        self.assertEqual(forbidden & set(sys.modules), set())


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------
class TestMarkdownExtraction(IngestionTestCase):
    def blocks(self):
        return si.ingest_source(self.markdown(), "src_md").blocks

    def test_heading_hierarchy_extracted(self) -> None:
        paths = [b.locator["heading_path"] for b in self.blocks()]
        self.assertEqual(
            paths,
            [
                ["Overview"],
                ["Overview", "Architecture"],
                ["Overview", "Architecture", "Validation"],
                ["Testing"],
                ["Testing", "Architecture"],
            ],
        )
        levels = [b.metadata["heading_level"] for b in self.blocks()]
        self.assertEqual(levels, [1, 2, 3, 1, 2])

    def test_repeated_headings_do_not_collide(self) -> None:
        """Two 'Architecture' sections under different parents stay distinct."""
        arch = [
            b for b in self.blocks() if b.metadata["heading_text"] == "Architecture"
        ]
        self.assertEqual(len(arch), 2)
        self.assertNotEqual(arch[0].locator["heading_path"], arch[1].locator["heading_path"])
        self.assertNotEqual(arch[0].locator["start_line"], arch[1].locator["start_line"])
        self.assertNotEqual(arch[0].block_id, arch[1].block_id)
        self.assertNotEqual(arch[0].locator["label"], arch[1].locator["label"])

    def test_line_ranges_are_one_based_and_ordered(self) -> None:
        lines = MARKDOWN_FIXTURE.split("\n")
        for block in self.blocks():
            start = block.locator["start_line"]
            end = block.locator["end_line"]
            self.assertGreaterEqual(start, 1)
            self.assertGreaterEqual(end, start)
            heading_text = block.metadata["heading_text"]
            if heading_text:
                # start_line must point at the actual heading line.
                self.assertIn(heading_text, lines[start - 1])
                self.assertTrue(lines[start - 1].lstrip().startswith("#"))

    def test_fenced_code_fake_heading_ignored(self) -> None:
        headings = [b.metadata["heading_text"] for b in self.blocks()]
        self.assertNotIn("Fake Heading", headings)
        for block in self.blocks():
            self.assertNotIn("Fake Heading", block.locator["label"])
        # The fence content is still carried inside its enclosing section.
        last = self.blocks()[-1]
        self.assertIn("# Fake Heading", last.text)

    def test_deterministic_ids_and_repeatability(self) -> None:
        path = self.markdown()
        first = si.ingest_source(path, "src_md").to_dict()
        second = si.ingest_source(path, "src_md").to_dict()
        self.assertEqual(first, second)
        for block in self.blocks():
            self.assertEqual(
                block.block_id,
                si.compute_block_id("src_md", block.locator, block.ordinal),
            )

    def test_utf8_and_traditional_chinese_preserved(self) -> None:
        content = "# \u6e2c\u8a66\u6587\u4ef6\n\n\u7e41\u9ad4\u4e2d\u6587\u5167\u5bb9\u3002\n"
        path = self.write("zh.md", content)
        blocks = si.ingest_source(path, "src_zh").blocks
        self.assertEqual(blocks[0].metadata["heading_text"], "\u6e2c\u8a66\u6587\u4ef6")
        self.assertIn("\u7e41\u9ad4\u4e2d\u6587\u5167\u5bb9\u3002", blocks[0].text)

    def test_preamble_before_first_heading_is_kept(self) -> None:
        path = self.write("pre.md", "Leading note.\n\n# Real Heading\n\nBody.\n")
        blocks = si.ingest_source(path, "src_pre").blocks
        self.assertEqual(blocks[0].locator["heading_path"], [])
        self.assertIn("Leading note.", blocks[0].text)
        self.assertEqual(blocks[1].metadata["heading_text"], "Real Heading")

    def test_markdown_without_headings_yields_one_block(self) -> None:
        path = self.write("flat.md", "Just prose, no headings at all.\n")
        blocks = si.ingest_source(path, "src_flat").blocks
        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0].locator["heading_path"], [])

    def test_crlf_and_lf_produce_identical_blocks(self) -> None:
        lf = self.write("lf.md", MARKDOWN_FIXTURE)
        crlf = self.write("crlf.md", MARKDOWN_FIXTURE.replace("\n", "\r\n").encode("utf-8"))
        lf_blocks = [b.to_dict() for b in si.ingest_source(lf, "src_x").blocks]
        crlf_blocks = [b.to_dict() for b in si.ingest_source(crlf, "src_x").blocks]
        self.assertEqual(lf_blocks, crlf_blocks)
        # ...but the raw digests differ, because the bytes genuinely differ.
        self.assertNotEqual(
            si.ingest_source(lf, "src_x").source_digest,
            si.ingest_source(crlf, "src_x").source_digest,
        )


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------
class TestTextExtraction(IngestionTestCase):
    def test_deterministic_paragraph_extraction(self) -> None:
        result = si.ingest_source(self.text(), "src_txt")
        self.assertEqual(result.block_count, 3)
        self.assertEqual([b.block_type for b in result.blocks], [si.BLOCK_PARAGRAPH] * 3)
        repeat = si.ingest_source(self.text("again.txt"), "src_txt")
        self.assertEqual(
            [b.block_id for b in result.blocks], [b.block_id for b in repeat.blocks]
        )

    def test_line_locators_are_one_based(self) -> None:
        blocks = si.ingest_source(self.text(), "src_txt").blocks
        ranges = [(b.locator["start"], b.locator["end"]) for b in blocks]
        self.assertEqual(ranges, [(1, 2), (4, 4), (7, 7)])
        for block in blocks:
            self.assertEqual(block.locator["kind"], "line_range")
            self.assertGreaterEqual(block.locator["start"], 1)

    def test_traditional_chinese_preserved(self) -> None:
        blocks = si.ingest_source(self.text(), "src_txt").blocks
        self.assertIn("\u7b2c\u4e00\u6bb5\u5167\u5bb9\u3002", blocks[0].text)
        self.assertIn("\u7b2c\u4e8c\u6bb5\u5167\u5bb9\u3002", blocks[1].text)

    def test_utf8_bom_handled(self) -> None:
        raw = b"\xef\xbb\xbf" + "\u7e41\u9ad4\u4e2d\u6587 with BOM.\n".encode("utf-8")
        path = self.write("bom.txt", raw)
        blocks = si.ingest_source(path, "src_bom").blocks
        self.assertEqual(len(blocks), 1)
        self.assertFalse(blocks[0].text.startswith("\ufeff"))
        self.assertTrue(blocks[0].text.startswith("\u7e41\u9ad4\u4e2d\u6587"))

    def test_unsupported_encoding_fails_deterministically(self) -> None:
        """Invalid UTF-8 must fail loudly rather than produce mojibake."""
        path = self.write("latin1.txt", b"caf\xe9 latin-1 bytes\n")
        with self.assertRaises(si.SourceEncodingUnsupported) as ctx:
            si.ingest_source(path, "src_latin1")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_ENCODING_UNSUPPORTED)

    def test_trailing_newline_does_not_shift_ranges(self) -> None:
        without = self.write("a.txt", "One paragraph.")
        with_nl = self.write("b.txt", "One paragraph.\n")
        a = si.ingest_source(without, "src_nl").blocks
        b = si.ingest_source(with_nl, "src_nl").blocks
        self.assertEqual(a[0].locator, b[0].locator)
        self.assertEqual(a[0].block_id, b[0].block_id)

    def test_empty_text_source_fails(self) -> None:
        path = self.write("empty.txt", "\n\n\n")
        with self.assertRaises(si.SourceTextUnavailable):
            si.ingest_source(path, "src_empty")


# ---------------------------------------------------------------------------
# Phase 12 compatibility (Phase 13 is an upstream producer, Phase 12 is frozen)
# ---------------------------------------------------------------------------
class TestPhase12Compatibility(IngestionTestCase):
    def test_source_digest_matches_phase12_definition(self) -> None:
        path = self.pdf()
        result = si.ingest_source(path, "src_pdf")
        self.assertEqual(result.source_digest, compute_source_digest(path.read_bytes()))
        # ...and is accepted verbatim by the Phase 12 inventory contract.
        inventory = SourceInventory.initialize(self.root, "proj")
        inventory.add_source("src_pdf", "pdf", source_digest=result.source_digest)
        inventory.save()
        self.assertFalse(inventory.source_changed("src_pdf", result.source_digest))
        self.assertTrue(inventory.source_changed("src_pdf", "0" * 64))

    def _assert_locator_accepted(self, path: Path, source_id: str, kind: str) -> None:
        result = si.ingest_source(path, source_id)
        inventory = SourceInventory.initialize(self.root, "proj")
        inventory.add_source(source_id, result.source_format, source_digest=result.source_digest)
        for block in result.blocks:
            locator = si.phase12_locator(block)
            self.assertEqual(locator["kind"], kind)
            # Frozen Phase 12 validator must accept the shape unchanged.
            self.assertEqual(validate_locator(locator), [])
            unit = inventory.add_unit(source_id, "generic", locator, "MEDIUM")
            self.assertTrue(unit["unit_id"].startswith(f"su:{source_id}:"))
        inventory.save()
        self.assertEqual(len(inventory.data["units"]), result.block_count)

    def test_pdf_locator_maps_into_phase12(self) -> None:
        self._assert_locator_accepted(self.pdf(), "src_pdf", "page")

    def test_markdown_locator_maps_into_phase12(self) -> None:
        self._assert_locator_accepted(self.markdown(), "src_md", "section")

    def test_text_locator_maps_into_phase12(self) -> None:
        self._assert_locator_accepted(self.text(), "src_txt", "line_range")

    def test_no_absolute_path_leaks_into_handoff(self) -> None:
        """Logical source identity must never depend on where the file lives."""
        for name, builder in (
            ("doc.pdf", lambda p: p.write_bytes(build_text_pdf(PAGE_TEXTS))),
            ("notes.md", lambda p: p.write_text(MARKDOWN_FIXTURE, encoding="utf-8")),
            ("plain.txt", lambda p: p.write_text(TEXT_FIXTURE, encoding="utf-8")),
        ):
            with tempfile.TemporaryDirectory() as d1, tempfile.TemporaryDirectory() as d2:
                p1, p2 = Path(d1) / name, Path(d2) / name
                builder(p1)
                builder(p2)
                r1 = si.ingest_source(p1, "src_same")
                r2 = si.ingest_source(p2, "src_same")
                self.assertEqual(r1.to_dict(), r2.to_dict(), name)
                serialized = json.dumps(r1.to_dict(), ensure_ascii=False)
                self.assertNotIn(d1, serialized)
                self.assertNotIn(d2, serialized)
                self.assertNotIn(str(Path.home()), serialized)

    def test_blocks_are_not_auto_promoted_to_source_units(self) -> None:
        """Extraction must not write any Phase 12 grounding artifact."""
        si.ingest_source(self.pdf(), "src_pdf")
        for artifact in (
            "source_inventory.json",
            "claim_traceability.json",
            "source_coverage.json",
            "source_grounded_qa.json",
        ):
            self.assertFalse((self.root / artifact).exists(), artifact)

    def test_extraction_result_carries_versions(self) -> None:
        result = si.ingest_source(self.markdown(), "src_md")
        self.assertEqual(result.schema_version, si.SCHEMA_VERSION)
        self.assertEqual(result.extractor_version, si.EXTRACTOR_VERSION)
        # Extractor version is independent of the release/tag version.
        self.assertNotIn("v", result.extractor_version)
        self.assertNotIn(".", result.extractor_version)


# ---------------------------------------------------------------------------
# DOCX (Phase 13.3)
# ---------------------------------------------------------------------------
class TestDocxExtraction(IngestionTestCase):
    def blocks(self):
        return si.ingest_source(self.docx(), "src_docx").blocks

    def test_paragraphs_extracted(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertIn("Synthetic governance overview.", texts)
        self.assertIn("Paragraph after the table.", texts)
        self.assertIn("Different repeated-heading section.", texts)

    def test_heading_level_1_hierarchy(self) -> None:
        top = [b for b in self.blocks() if b.locator["heading_path"][:1] == ["Governance"]]
        self.assertTrue(top)
        first = self.blocks()[0]
        self.assertEqual(first.locator["heading_path"], ["Governance"])
        self.assertEqual(first.metadata["heading_level"], 1)

    def test_heading_level_2_hierarchy(self) -> None:
        nested = [
            b
            for b in self.blocks()
            if b.locator["heading_path"] == ["Governance", "Risk Controls"]
            and b.block_type == si.BLOCK_DOCX_SECTION
        ]
        self.assertTrue(nested)
        self.assertEqual(nested[0].metadata["heading_level"], 2)

    def test_repeated_heading_labels_remain_distinct(self) -> None:
        """'Risk Controls' under two different H1 parents must not collide."""
        sections = [
            b
            for b in self.blocks()
            if b.block_type == si.BLOCK_DOCX_SECTION
            and b.metadata.get("heading_text") == "Risk Controls"
            and b.metadata.get("heading_level") == 2
        ]
        paths = [b.locator["heading_path"] for b in sections]
        self.assertIn(["Governance", "Risk Controls"], paths)
        self.assertIn(["Testing", "Risk Controls"], paths)
        labels = {b.locator["label"] for b in sections}
        self.assertIn("Governance > Risk Controls", labels)
        self.assertIn("Testing > Risk Controls", labels)
        ids = [b.block_id for b in sections]
        self.assertEqual(len(ids), len(set(ids)))

    def test_traditional_chinese_and_english_preserved(self) -> None:
        path = self.write("bilingual.docx", build_bilingual_docx())
        texts = "\n".join(b.text for b in si.ingest_source(path, "src_bi").blocks)
        self.assertIn("\u6cbb\u7406\u67b6\u69cb\u8207\u98a8\u96aa\u63a7\u5236\u3002", texts)
        self.assertIn("Governance and Risk Controls", texts)
        self.assertIn("\u6e2c\u8a66\u6587\u4ef6", texts)

    def test_table_extracted(self) -> None:
        tables = [b for b in self.blocks() if b.block_type == si.BLOCK_DOCX_TABLE]
        self.assertEqual(len(tables), 1)
        table = tables[0]
        self.assertEqual(table.locator["table_index"], 1)
        self.assertEqual(table.locator["start_row"], 1)
        self.assertEqual(table.locator["end_row"], 3)
        self.assertEqual(table.metadata["row_count"], 3)
        self.assertEqual(table.metadata["column_count"], 2)
        for cell in ("Control", "Owner", "Access", "Security", "Review", "Governance"):
            self.assertIn(cell, table.text)

    def test_table_only_document_is_not_discarded(self) -> None:
        path = self.write("tableonly.docx", build_table_only_docx())
        blocks = si.ingest_source(path, "src_tbl").blocks
        self.assertEqual([b.block_type for b in blocks], [si.BLOCK_DOCX_TABLE])
        self.assertIn("\u9805\u76ee", blocks[0].text)

    def test_table_reading_order_is_row_major(self) -> None:
        table = next(b for b in self.blocks() if b.block_type == si.BLOCK_DOCX_TABLE)
        rows = table.text.split("\n")
        self.assertEqual(rows, ["Control | Owner", "Access | Security", "Review | Governance"])
        # Row order top-to-bottom, cell order left-to-right within each row.
        self.assertLess(rows.index("Control | Owner"), rows.index("Access | Security"))

    def test_merged_cell_behaviour_is_documented_and_stable(self) -> None:
        """python-docx repeats a merged cell across the positions it spans."""
        from docx import Document

        import io as _io

        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        for r in range(2):
            for c in range(3):
                table.cell(r, c).text = f"r{r + 1}c{c + 1}"
        table.cell(0, 0).merge(table.cell(0, 1))
        buffer = _io.BytesIO()
        doc.save(buffer)
        path = self.write("merged.docx", buffer.getvalue())

        block = si.ingest_source(path, "src_merged").blocks[0]
        rows = block.text.split("\n")
        # One output line per table row, even though a merged cell's own text
        # contains a line break.
        self.assertEqual(len(rows), 2)
        # The merged text is present and repeated, never dropped.
        self.assertEqual(rows[0].count("r1c1"), 2)
        self.assertEqual(rows[0].count("r1c2"), 2)
        self.assertIn("r1c3", rows[0])
        self.assertEqual(rows[1], "r2c1 | r2c2 | r2c3")
        self.assertEqual(block.metadata["row_count"], 2)

    def test_paragraph_and_table_source_order_preserved(self) -> None:
        """A table must appear between the paragraphs that surround it."""
        blocks = self.blocks()
        kinds = [b.block_type for b in blocks]
        table_at = kinds.index(si.BLOCK_DOCX_TABLE)
        before = "\n".join(b.text for b in blocks[:table_at])
        after = "\n".join(b.text for b in blocks[table_at + 1 :])
        self.assertIn("Synthetic governance overview.", before)
        self.assertIn("Paragraph after the table.", after)
        self.assertNotIn("Paragraph after the table.", before)
        # Ordinals stay strictly increasing in document order.
        ordinals = [b.ordinal for b in blocks]
        self.assertEqual(ordinals, list(range(1, len(blocks) + 1)))

    def test_structural_element_ranges_are_one_based(self) -> None:
        for block in self.blocks():
            if block.block_type == si.BLOCK_DOCX_SECTION:
                self.assertGreaterEqual(block.locator["start_element"], 1)
                self.assertGreaterEqual(
                    block.locator["end_element"], block.locator["start_element"]
                )
            else:
                self.assertGreaterEqual(block.locator["start_row"], 1)
                self.assertGreaterEqual(block.locator["table_index"], 1)

    def test_no_page_locator_is_fabricated(self) -> None:
        """DOCX is flow-based: no rendered page numbers may be invented."""
        for block in self.blocks():
            self.assertEqual(block.locator["kind"], "section")
            self.assertNotIn("page", block.locator)
            self.assertNotIn("page", block.metadata)
            self.assertNotEqual(block.block_type, si.BLOCK_PAGE)

    def test_stable_block_ids_and_repeatability(self) -> None:
        path = self.docx()
        first = si.ingest_source(path, "src_docx").to_dict()
        second = si.ingest_source(path, "src_docx").to_dict()
        self.assertEqual(first, second)
        for block in si.ingest_source(path, "src_docx").blocks:
            self.assertEqual(
                block.block_id,
                si.compute_block_id("src_docx", block.locator, block.ordinal),
            )
            self.assertTrue(block.block_id.startswith("blk:src_docx:"))

    def test_unstyled_bold_text_is_not_a_heading(self) -> None:
        """Headings come from styles only, never from font weight or size."""
        path = self.write("bold.docx", build_unstyled_bold_docx())
        blocks = si.ingest_source(path, "src_bold").blocks
        for block in blocks:
            self.assertEqual(block.locator["heading_path"], [])
            self.assertEqual(block.metadata.get("heading_level"), 0)
        self.assertIn("Looks Like A Heading", blocks[0].text)

    def test_empty_docx_fails_with_text_unavailable(self) -> None:
        path = self.write("empty.docx", build_empty_docx())
        self.assertEqual(si.detect_source_format(path), si.FORMAT_DOCX)
        with self.assertRaises(si.SourceTextUnavailable) as ctx:
            si.ingest_source(path, "src_empty_docx")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_TEXT_UNAVAILABLE)

    def test_corrupt_docx_fails_with_extraction_failed(self) -> None:
        path = self.write("corrupt.docx", build_corrupt_docx())
        with self.assertRaises(si.SourceExtractionFailed) as ctx:
            si.ingest_source(path, "src_corrupt_docx")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_EXTRACTION_FAILED)
        self.assertNotEqual(
            ctx.exception.error_code, si.ERROR_SOURCE_TEXT_UNAVAILABLE
        )

    def test_encrypted_docx_fails_deterministically(self) -> None:
        """No interactive password prompt, no password handling."""
        path = self.write("locked.docx", build_encrypted_like_docx())
        with self.assertRaises(si.SourceExtractionFailed) as ctx:
            si.ingest_source(path, "src_locked")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_EXTRACTION_FAILED)
        self.assertIn("encrypted", str(ctx.exception).lower())

    def test_image_only_docx_does_not_trigger_ocr(self) -> None:
        path = self.write("imageonly.docx", build_image_only_docx())
        with self.assertRaises(si.SourceTextUnavailable) as ctx:
            si.ingest_source(path, "src_imageonly")
        self.assertIn("OCR is not supported", str(ctx.exception))
        forbidden = {"pytesseract", "tesserocr", "easyocr", "paddleocr"}
        self.assertEqual(forbidden & set(sys.modules), set())

    def test_layout_only_empty_paragraphs_produce_no_blocks(self) -> None:
        from docx import Document

        import io as _io

        doc = Document()
        doc.add_paragraph("   ")
        doc.add_paragraph("")
        doc.add_paragraph("Real content.")
        doc.add_paragraph("\t")
        buffer = _io.BytesIO()
        doc.save(buffer)
        path = self.write("sparse.docx", buffer.getvalue())

        blocks = si.ingest_source(path, "src_sparse").blocks
        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0].text, "Real content.")

    def test_no_phase12_artifact_written(self) -> None:
        si.ingest_source(self.docx(), "src_docx")
        for artifact in (
            "source_inventory.json",
            "claim_traceability.json",
            "source_coverage.json",
            "source_grounded_qa.json",
        ):
            self.assertFalse((self.root / artifact).exists(), artifact)

    def test_absolute_path_independence(self) -> None:
        payload = build_structured_docx()
        with tempfile.TemporaryDirectory() as d1, tempfile.TemporaryDirectory() as d2:
            p1, p2 = Path(d1) / "a.docx", Path(d2) / "a.docx"
            p1.write_bytes(payload)
            p2.write_bytes(payload)
            r1 = si.ingest_source(p1, "src_same")
            r2 = si.ingest_source(p2, "src_same")
            self.assertEqual(r1.to_dict(), r2.to_dict())
            serialized = json.dumps(r1.to_dict(), ensure_ascii=False)
            self.assertNotIn(d1, serialized)
            self.assertNotIn(d2, serialized)

    def test_source_digest_is_raw_bytes(self) -> None:
        """DOCX must not introduce a competing unzipped/normalized digest."""
        path = self.docx()
        result = si.ingest_source(path, "src_docx")
        self.assertEqual(result.source_digest, compute_source_digest(path.read_bytes()))

    def test_docx_locator_accepted_by_frozen_phase12_validator(self) -> None:
        result = si.ingest_source(self.docx(), "src_docx")
        inventory = SourceInventory.initialize(self.root, "proj")
        inventory.add_source("src_docx", "docx", source_digest=result.source_digest)
        for block in result.blocks:
            locator = si.phase12_locator(block)
            self.assertEqual(locator["kind"], "section")
            self.assertEqual(validate_locator(locator), [])
            unit = inventory.add_unit("src_docx", "generic", locator, "MEDIUM")
            self.assertTrue(unit["unit_id"].startswith("su:src_docx:"))
        inventory.save()
        self.assertEqual(len(inventory.data["units"]), result.block_count)

    def test_no_priority_is_assigned_by_extraction(self) -> None:
        """Heading level is structure, never a semantic priority."""
        for block in self.blocks():
            serialized = json.dumps(block.to_dict(), ensure_ascii=False)
            for word in ("HIGH", "MEDIUM", "LOW"):
                self.assertNotIn(word, serialized)


# ---------------------------------------------------------------------------
# HTML (Phase 13.4)
# ---------------------------------------------------------------------------
class TestHtmlExtraction(IngestionTestCase):
    def blocks(self):
        return si.ingest_source(self.html(), "src_html").blocks

    def test_heading_hierarchy_extracted(self) -> None:
        sections = [b for b in self.blocks() if b.block_type == si.BLOCK_HTML_SECTION]
        paths = [b.locator["heading_path"] for b in sections]
        self.assertIn(["Architecture"], paths)
        self.assertIn(["Architecture", "Overview"], paths)
        self.assertIn(["Testing"], paths)
        first = sections[0]
        self.assertEqual(first.metadata["heading_level"], 1)

    def test_repeated_headings_remain_distinct(self) -> None:
        """'Overview' under two different h1 parents must not collide.

        Note a heading's content can span more than one block: when a list or
        table interrupts, the following paragraphs form a further block under the
        same heading path, which is what preserves document order.
        """
        overviews = [
            b
            for b in self.blocks()
            if b.metadata.get("heading_text") == "Overview"
            and b.block_type == si.BLOCK_HTML_SECTION
        ]
        paths = {tuple(b.locator["heading_path"]) for b in overviews}
        self.assertEqual(
            paths, {("Architecture", "Overview"), ("Testing", "Overview")}
        )
        labels = {b.locator["label"] for b in overviews}
        self.assertIn("Architecture > Overview", labels)
        self.assertIn("Testing > Overview", labels)
        ids = [b.block_id for b in overviews]
        self.assertEqual(len(ids), len(set(ids)), "block ids must stay unique")
        # The two headings' own sections carry different text.
        arch = next(
            b for b in overviews if b.locator["heading_path"] == ["Architecture", "Overview"]
        )
        test = next(
            b for b in overviews if b.locator["heading_path"] == ["Testing", "Overview"]
        )
        self.assertNotEqual(arch.block_id, test.block_id)
        self.assertIn("Governance Architecture", arch.text)
        self.assertIn("second overview", test.text)

    def test_heading_level_jump_invents_nothing(self) -> None:
        """h1 -> h3 nests one deeper; no missing h2 is fabricated."""
        deep = next(
            b for b in self.blocks() if b.metadata.get("heading_text") == "Deep Jump"
        )
        self.assertEqual(deep.locator["heading_path"], ["Testing", "Deep Jump"])
        self.assertEqual(deep.metadata["heading_level"], 3)

    def test_paragraphs_extracted(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertIn("Intro with bold, emphasis and Example Link.", texts)
        self.assertIn("After table", texts)
        self.assertIn("second overview", texts)

    def test_traditional_chinese_and_english_preserved(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertIn("\u6cbb\u7406\u67b6\u69cb", texts)
        self.assertIn("Governance Architecture", texts)

    def test_unordered_list_extracted(self) -> None:
        lists = [b for b in self.blocks() if b.block_type == si.BLOCK_HTML_LIST]
        unordered = next(b for b in lists if b.metadata["list_type"] == "unordered")
        self.assertIn("- alpha", unordered.text)
        self.assertIn("- beta", unordered.text)

    def test_ordered_list_preserves_numbering(self) -> None:
        lists = [b for b in self.blocks() if b.block_type == si.BLOCK_HTML_LIST]
        ordered = next(b for b in lists if b.metadata["list_type"] == "ordered")
        self.assertEqual(ordered.text, "1. first\n2. second")

    def test_nested_list_is_deterministic_and_not_duplicated(self) -> None:
        lists = [b for b in self.blocks() if b.block_type == si.BLOCK_HTML_LIST]
        unordered = next(b for b in lists if b.metadata["list_type"] == "unordered")
        self.assertEqual(unordered.text, "- alpha\n- beta\n  - beta-nested")
        # The nested list must not also appear as its own block.
        all_text = "\n".join(b.text for b in self.blocks())
        self.assertEqual(all_text.count("beta-nested"), 1)

    def test_table_extracted(self) -> None:
        table = next(b for b in self.blocks() if b.block_type == si.BLOCK_HTML_TABLE)
        self.assertEqual(table.locator["table_index"], 1)
        self.assertEqual(table.locator["start_row"], 1)
        self.assertEqual(table.locator["end_row"], 4)
        for cell in ("Control", "Owner", "Access", "Security", "Review", "Wide cell"):
            self.assertIn(cell, table.text)

    def test_table_row_order_is_deterministic(self) -> None:
        table = next(b for b in self.blocks() if b.block_type == si.BLOCK_HTML_TABLE)
        rows = table.text.split("\n")
        self.assertEqual(rows[0], "Control | Owner")
        self.assertEqual(rows[1], "Access | Security")
        self.assertEqual(len(rows), 4)

    def test_rowspan_colspan_are_not_expanded_into_a_grid(self) -> None:
        """DOM cell order is reported; no visual spreadsheet model is invented."""
        table = next(b for b in self.blocks() if b.block_type == si.BLOCK_HTML_TABLE)
        rows = table.text.split("\n")
        # The rowspan cell appears once, on the row where it is declared.
        self.assertEqual(rows[1], "Access | Security")
        self.assertEqual(rows[2], "Review")
        self.assertEqual(table.text.count("Access"), 1)
        # The colspan cell is a single cell, not duplicated across columns.
        self.assertEqual(rows[3], "Wide cell")
        self.assertEqual(table.metadata["row_count"], 4)

    def test_mixed_document_order_preserved(self) -> None:
        blocks = self.blocks()
        kinds = [b.block_type for b in blocks]
        table_at = kinds.index(si.BLOCK_HTML_TABLE)
        first_list_at = kinds.index(si.BLOCK_HTML_LIST)
        before = "\n".join(b.text for b in blocks[:table_at])
        after = "\n".join(b.text for b in blocks[table_at + 1 :])
        self.assertLess(first_list_at, table_at, "lists precede the table in source")
        self.assertIn("Governance Architecture", before)
        self.assertIn("After table", after)
        self.assertNotIn("After table", before)
        self.assertEqual([b.ordinal for b in blocks], list(range(1, len(blocks) + 1)))

    def test_visible_hyperlink_text_retained(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertIn("Example Link", texts)

    def test_hyperlink_target_is_not_source_evidence(self) -> None:
        """Link text is kept; the destination is never followed or ingested."""
        serialized = json.dumps(
            si.ingest_source(self.html(), "src_html").to_dict(), ensure_ascii=False
        )
        self.assertNotIn("example.invalid/page", serialized)

    def test_script_content_excluded(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertNotIn("var a = 1", texts)
        self.assertNotIn("app.js", texts)

    def test_style_content_excluded(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertNotIn("display:none", texts)
        self.assertNotIn(".hidden", texts)

    def test_comments_excluded(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertNotIn("must never be ingested", texts)

    def test_json_ld_excluded(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertNotIn("JSONLD Headline", texts)
        self.assertNotIn("@type", texts)

    def test_noscript_not_ingested(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertNotIn("noscript content", texts)

    def test_title_and_head_not_ingested(self) -> None:
        texts = "\n".join(b.text for b in self.blocks())
        self.assertNotIn("Ignored Title", texts)

    def test_iframe_and_image_references_not_resolved(self) -> None:
        serialized = json.dumps(
            si.ingest_source(self.html(), "src_html").to_dict(), ensure_ascii=False
        )
        for reference in ("example.invalid/frame", "example.invalid/image.png"):
            self.assertNotIn(reference, serialized)

    def test_empty_html_fails_with_text_unavailable(self) -> None:
        path = self.write("empty.html", "<html><head></head><body></body></html>\n")
        self.assertEqual(si.detect_source_format(path), si.FORMAT_HTML)
        with self.assertRaises(si.SourceTextUnavailable) as ctx:
            si.ingest_source(path, "src_emptyhtml")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_TEXT_UNAVAILABLE)

    def test_script_only_html_fails_without_javascript(self) -> None:
        path = self.write(
            "scriptonly.html",
            '<html><body><script>document.write("content")</script></body></html>\n',
        )
        with self.assertRaises(si.SourceTextUnavailable) as ctx:
            si.ingest_source(path, "src_scriptonly")
        self.assertIn("JavaScript is not executed", str(ctx.exception))

    def test_malformed_recoverable_html_is_deterministic(self) -> None:
        path = self.write(
            "malformed.htm", "<html><body><p>unclosed<div>mixed</p></body>\n"
        )
        first = si.ingest_source(path, "src_bad").to_dict()
        second = si.ingest_source(path, "src_bad").to_dict()
        self.assertEqual(first, second)
        self.assertIn("unclosed", first["blocks"][0]["text"])

    def test_stable_block_ids_and_repeatability(self) -> None:
        path = self.html()
        first = si.ingest_source(path, "src_html").to_dict()
        second = si.ingest_source(path, "src_html").to_dict()
        self.assertEqual(first, second)
        for block in si.ingest_source(path, "src_html").blocks:
            self.assertEqual(
                block.block_id,
                si.compute_block_id("src_html", block.locator, block.ordinal),
            )
            self.assertTrue(block.block_id.startswith("blk:src_html:"))

    def test_structural_ordinals_are_one_based(self) -> None:
        for block in self.blocks():
            self.assertGreaterEqual(block.locator["start_element"], 1)
            self.assertGreaterEqual(
                block.locator["end_element"], block.locator["start_element"]
            )

    def test_no_page_locator_is_fabricated(self) -> None:
        for block in self.blocks():
            self.assertEqual(block.locator["kind"], "section")
            for forbidden in ("page", "scroll", "pixel", "offset"):
                self.assertNotIn(forbidden, block.locator)
                self.assertNotIn(forbidden, block.metadata)

    def test_absolute_path_independence(self) -> None:
        with tempfile.TemporaryDirectory() as d1, tempfile.TemporaryDirectory() as d2:
            p1, p2 = Path(d1) / "a.html", Path(d2) / "a.html"
            p1.write_text(HTML_FIXTURE, encoding="utf-8")
            p2.write_text(HTML_FIXTURE, encoding="utf-8")
            r1 = si.ingest_source(p1, "src_same")
            r2 = si.ingest_source(p2, "src_same")
            self.assertEqual(r1.to_dict(), r2.to_dict())
            serialized = json.dumps(r1.to_dict(), ensure_ascii=False)
            self.assertNotIn(d1, serialized)
            self.assertNotIn(d2, serialized)

    def test_source_digest_is_raw_bytes(self) -> None:
        path = self.html()
        result = si.ingest_source(path, "src_html")
        self.assertEqual(result.source_digest, compute_source_digest(path.read_bytes()))

    def test_equivalent_dom_different_bytes_may_differ_in_digest(self) -> None:
        """Whitespace-only markup differences keep Phase 12 change semantics."""
        tight = self.write("a.html", "<html><body><p>Hello</p></body></html>")
        loose = self.write("b.html", "<html><body><p>\n  Hello\n</p></body></html>")
        a = si.ingest_source(tight, "src_eq")
        b = si.ingest_source(loose, "src_eq")
        self.assertEqual(a.blocks[0].text, b.blocks[0].text)
        self.assertNotEqual(a.source_digest, b.source_digest)

    def test_utf8_bom_handled(self) -> None:
        raw = b"\xef\xbb\xbf" + HTML_FIXTURE.encode("utf-8")
        path = self.write("bom.html", raw)
        blocks = si.ingest_source(path, "src_bomhtml").blocks
        self.assertIn("\u6cbb\u7406\u67b6\u69cb", "\n".join(b.text for b in blocks))

    def test_unsupported_encoding_fails_clearly(self) -> None:
        path = self.write("latin1.html", b"<html><body><p>caf\xe9</p></body></html>")
        with self.assertRaises(si.SourceEncodingUnsupported) as ctx:
            si.ingest_source(path, "src_latinhtml")
        self.assertEqual(ctx.exception.error_code, si.ERROR_SOURCE_ENCODING_UNSUPPORTED)

    def test_html_locator_accepted_by_frozen_phase12_validator(self) -> None:
        result = si.ingest_source(self.html(), "src_html")
        inventory = SourceInventory.initialize(self.root, "proj")
        inventory.add_source("src_html", "html", source_digest=result.source_digest)
        for block in result.blocks:
            locator = si.phase12_locator(block)
            self.assertEqual(locator["kind"], "section")
            self.assertEqual(validate_locator(locator), [])
            unit = inventory.add_unit("src_html", "generic", locator, "MEDIUM")
            self.assertTrue(unit["unit_id"].startswith("su:src_html:"))
        inventory.save()
        self.assertEqual(len(inventory.data["units"]), result.block_count)

    def test_no_phase12_artifact_written(self) -> None:
        si.ingest_source(self.html(), "src_html")
        for artifact in (
            "source_inventory.json",
            "claim_traceability.json",
            "source_coverage.json",
            "source_grounded_qa.json",
        ):
            self.assertFalse((self.root / artifact).exists(), artifact)

    def test_no_priority_is_assigned_by_extraction(self) -> None:
        for block in self.blocks():
            serialized = json.dumps(block.to_dict(), ensure_ascii=False)
            for word in ("HIGH", "MEDIUM", "LOW"):
                self.assertNotIn(word, serialized)


# ---------------------------------------------------------------------------
# HTML network-safety contract (Phase 13.4)
# ---------------------------------------------------------------------------
class TestHtmlNetworkSafety(IngestionTestCase):
    """Prove no network access, rather than relying on the network being down."""

    def test_no_socket_is_opened_during_html_ingestion(self) -> None:
        import socket

        calls: list[str] = []

        def deny(*args: object, **kwargs: object):
            calls.append("socket")
            raise AssertionError("HTML ingestion attempted to open a socket")

        originals = {
            "socket": socket.socket,
            "create_connection": socket.create_connection,
            "getaddrinfo": socket.getaddrinfo,
        }
        socket.socket = deny  # type: ignore[assignment]
        socket.create_connection = deny  # type: ignore[assignment]
        socket.getaddrinfo = deny  # type: ignore[assignment]
        try:
            result = si.ingest_source(self.html(), "src_nonet")
        finally:
            socket.socket = originals["socket"]  # type: ignore[assignment]
            socket.create_connection = originals["create_connection"]  # type: ignore[assignment]
            socket.getaddrinfo = originals["getaddrinfo"]  # type: ignore[assignment]

        self.assertEqual(calls, [])
        self.assertGreater(result.block_count, 0)

    def test_no_urlopen_during_html_ingestion(self) -> None:
        import urllib.request

        calls: list[str] = []

        def deny(*args: object, **kwargs: object):
            calls.append("urlopen")
            raise AssertionError("HTML ingestion attempted an HTTP request")

        original = urllib.request.urlopen
        urllib.request.urlopen = deny  # type: ignore[assignment]
        try:
            si.ingest_source(self.html(), "src_nonet2")
        finally:
            urllib.request.urlopen = original  # type: ignore[assignment]
        self.assertEqual(calls, [])

    def test_local_relative_assets_are_not_read(self) -> None:
        """A referenced local file must never be opened by the parser."""
        secret = self.root / "secret.txt"
        secret.write_text("SHOULD-NEVER-BE-READ", encoding="utf-8")
        sibling = self.root / "other.html"
        sibling.write_text(
            "<html><body><p>SIBLING-CONTENT</p></body></html>", encoding="utf-8"
        )
        page = self.write(
            "page.html",
            "<html><body>"
            '<h1>Main</h1><p>Only this counts.</p>'
            '<img src="secret.txt">'
            '<iframe src="other.html"></iframe>'
            '<link rel="stylesheet" href="secret.txt">'
            "</body></html>",
        )
        serialized = json.dumps(
            si.ingest_source(page, "src_local").to_dict(), ensure_ascii=False
        )
        self.assertIn("Only this counts.", serialized)
        self.assertNotIn("SHOULD-NEVER-BE-READ", serialized)
        self.assertNotIn("SIBLING-CONTENT", serialized)

    def test_no_browser_or_ocr_module_is_imported(self) -> None:
        si.ingest_source(self.html(), "src_nobrowser")
        forbidden = {
            "playwright",
            "selenium",
            "pyppeteer",
            "pytesseract",
            "tesserocr",
            "easyocr",
        }
        self.assertEqual(forbidden & set(sys.modules), set())


# ---------------------------------------------------------------------------
# Cross-format regression: Phase 13.3 must not change earlier extractors
# ---------------------------------------------------------------------------
class TestExistingExtractorsUnchanged(IngestionTestCase):
    def test_pdf_still_uses_one_based_page_locators(self) -> None:
        result = si.ingest_source(self.pdf(), "src_pdf")
        self.assertEqual([b.locator["start"] for b in result.blocks], [1, 2, 3])
        self.assertEqual([b.block_type for b in result.blocks], [si.BLOCK_PAGE] * 3)

    def test_markdown_still_produces_section_blocks(self) -> None:
        blocks = si.ingest_source(self.markdown(), "src_md").blocks
        self.assertEqual(len(blocks), 5)
        self.assertEqual(
            [b.block_type for b in blocks], [si.BLOCK_MARKDOWN_SECTION] * 5
        )
        self.assertEqual(blocks[0].locator["heading_path"], ["Overview"])

    def test_text_still_produces_paragraph_blocks(self) -> None:
        blocks = si.ingest_source(self.text(), "src_txt").blocks
        self.assertEqual([b.block_type for b in blocks], [si.BLOCK_PARAGRAPH] * 3)
        self.assertEqual(
            [(b.locator["start"], b.locator["end"]) for b in blocks],
            [(1, 2), (4, 4), (7, 7)],
        )

    def test_detection_unchanged_for_earlier_formats(self) -> None:
        self.assertEqual(si.detect_source_format(self.pdf()), si.FORMAT_PDF)
        self.assertEqual(si.detect_source_format(self.markdown()), si.FORMAT_MARKDOWN)
        self.assertEqual(si.detect_source_format(self.text()), si.FORMAT_TEXT)
        self.assertEqual(si.detect_source_format(self.docx()), si.FORMAT_DOCX)
        fake = self.write("fake.pdf", b"this is not a pdf at all")
        self.assertEqual(si.detect_source_format(fake), si.FORMAT_UNSUPPORTED)

    def test_docx_behaviour_unchanged(self) -> None:
        """Phase 13.4 must not alter the DOCX block or locator shape."""
        blocks = si.ingest_source(self.docx(), "src_docx").blocks
        self.assertEqual(len(blocks), 6)
        self.assertEqual(
            [b.block_type for b in blocks],
            [
                si.BLOCK_DOCX_SECTION,
                si.BLOCK_DOCX_SECTION,
                si.BLOCK_DOCX_TABLE,
                si.BLOCK_DOCX_SECTION,
                si.BLOCK_DOCX_SECTION,
                si.BLOCK_DOCX_SECTION,
            ],
        )
        table = blocks[2]
        self.assertEqual(table.locator["table_index"], 1)
        self.assertEqual(table.locator["end_row"], 3)
        # DOCX table locators intentionally keep their existing shape.
        self.assertNotIn("start_element", table.locator)
        self.assertEqual(blocks[0].locator["heading_path"], ["Governance"])

    def test_extractor_version_unchanged_shape(self) -> None:
        for path, source_id in (
            (self.pdf(), "src_pdf"),
            (self.markdown(), "src_md"),
            (self.text(), "src_txt"),
            (self.docx(), "src_docx"),
            (self.html(), "src_html"),
        ):
            result = si.ingest_source(path, source_id)
            self.assertEqual(result.extractor_version, si.EXTRACTOR_VERSION)
            self.assertEqual(result.schema_version, si.SCHEMA_VERSION)


# ---------------------------------------------------------------------------
# CLI adapter
# ---------------------------------------------------------------------------
class TestCliAdapter(IngestionTestCase):
    def run_cli(self, argv: list[str]) -> tuple[int, str, str]:
        import contextlib
        import io

        import ingest_source as cli

        out, err = io.StringIO(), io.StringIO()
        with contextlib.redirect_stdout(out), contextlib.redirect_stderr(err):
            code = cli.main(argv)
        return code, out.getvalue(), err.getvalue()

    def test_cli_writes_output_and_summary(self) -> None:
        out_path = self.root / "out" / "result.json"
        code, stdout, _ = self.run_cli(
            [
                "--source",
                str(self.pdf()),
                "--source-id",
                "src_cli",
                "--output",
                str(out_path),
            ]
        )
        self.assertEqual(code, 0)
        summary = json.loads(stdout)
        self.assertEqual(summary["source_id"], "src_cli")
        self.assertEqual(summary["source_format"], "pdf")
        self.assertEqual(summary["block_count"], 3)
        self.assertEqual(len(summary["source_digest"]), 64)
        written = json.loads(out_path.read_text(encoding="utf-8"))
        self.assertEqual(len(written["blocks"]), 3)

    def test_cli_stdout_json_when_no_output(self) -> None:
        code, stdout, _ = self.run_cli(
            ["--source", str(self.markdown()), "--source-id", "src_cli"]
        )
        self.assertEqual(code, 0)
        self.assertEqual(len(json.loads(stdout)["blocks"]), 5)

    def test_cli_detect_only(self) -> None:
        code, stdout, _ = self.run_cli(
            ["--source", str(self.text()), "--source-id", "src_cli", "--detect-only"]
        )
        self.assertEqual(code, 0)
        self.assertEqual(json.loads(stdout)["source_format"], "text")

    def test_cli_error_is_concise_with_stable_code(self) -> None:
        scan = self.write("scan.pdf", build_textless_pdf(1))
        code, stdout, stderr = self.run_cli(
            ["--source", str(scan), "--source-id", "src_scan"]
        )
        self.assertEqual(code, 1)
        self.assertEqual(stdout, "")
        self.assertTrue(stderr.startswith(si.ERROR_SOURCE_TEXT_UNAVAILABLE))
        self.assertNotIn("Traceback", stderr)

    def test_cli_unsupported_format_error(self) -> None:
        docx = self.write("deck.docx", b"PK\x03\x04 not really a docx")
        code, _, stderr = self.run_cli(
            ["--source", str(docx), "--source-id", "src_docx"]
        )
        self.assertEqual(code, 1)
        self.assertTrue(stderr.startswith(si.ERROR_SOURCE_EXTRACTION_FAILED))

    def test_cli_handles_docx_through_same_api(self) -> None:
        out_path = self.root / "docx.json"
        code, stdout, _ = self.run_cli(
            [
                "--source",
                str(self.docx()),
                "--source-id",
                "src_docx",
                "--output",
                str(out_path),
            ]
        )
        self.assertEqual(code, 0)
        summary = json.loads(stdout)
        self.assertEqual(summary["source_format"], "docx")
        self.assertEqual(summary["block_count"], 6)
        written = json.loads(out_path.read_text(encoding="utf-8"))
        self.assertEqual(len(written["blocks"]), 6)

    def test_cli_docx_detect_only(self) -> None:
        code, stdout, _ = self.run_cli(
            ["--source", str(self.docx()), "--source-id", "src_docx", "--detect-only"]
        )
        self.assertEqual(code, 0)
        self.assertEqual(json.loads(stdout)["source_format"], "docx")

    def test_cli_docx_error_has_no_traceback(self) -> None:
        path = self.write("empty.docx", build_empty_docx())
        code, stdout, stderr = self.run_cli(
            ["--source", str(path), "--source-id", "src_empty"]
        )
        self.assertEqual(code, 1)
        self.assertEqual(stdout, "")
        self.assertTrue(stderr.startswith(si.ERROR_SOURCE_TEXT_UNAVAILABLE))
        self.assertNotIn("Traceback", stderr)

    def test_cli_handles_html_through_same_api(self) -> None:
        out_path = self.root / "html.json"
        code, stdout, _ = self.run_cli(
            [
                "--source",
                str(self.html()),
                "--source-id",
                "src_html",
                "--output",
                str(out_path),
            ]
        )
        self.assertEqual(code, 0)
        summary = json.loads(stdout)
        self.assertEqual(summary["source_format"], "html")
        self.assertEqual(summary["block_count"], 9)
        written = json.loads(out_path.read_text(encoding="utf-8"))
        self.assertEqual(len(written["blocks"]), 9)

    def test_cli_has_no_url_option(self) -> None:
        """Phase 13.4 is local-file only: no --url input exists."""
        import ingest_source as cli

        actions = {a.option_strings[0] for a in cli.build_parser()._actions if a.option_strings}
        self.assertNotIn("--url", actions)
        self.assertIn("--source", actions)

    def test_cli_html_error_has_no_traceback(self) -> None:
        path = self.write("empty.html", "<html><head></head><body></body></html>")
        code, stdout, stderr = self.run_cli(
            ["--source", str(path), "--source-id", "src_eh"]
        )
        self.assertEqual(code, 1)
        self.assertEqual(stdout, "")
        self.assertTrue(stderr.startswith(si.ERROR_SOURCE_TEXT_UNAVAILABLE))
        self.assertNotIn("Traceback", stderr)


if __name__ == "__main__":
    unittest.main(verbosity=2)
